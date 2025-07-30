import pdfplumber
import docx
from typing import Dict, Any, List
import re
from io import BytesIO
import logging

logger = logging.getLogger(__name__)

class ResumeParser:
    """Service for parsing resume files (PDF, DOCX) and extracting structured data"""
    
    def __init__(self):
        self.section_patterns = {
            'contact': [
                r'contact\s+information',
                r'personal\s+information',
                r'contact\s+details'
            ],
            'summary': [
                r'professional\s+summary',
                r'career\s+summary',
                r'summary',
                r'profile',
                r'objective',
                r'career\s+objective'
            ],
            'experience': [
                r'work\s+experience',
                r'professional\s+experience',
                r'employment\s+history',
                r'experience',
                r'career\s+history'
            ],
            'education': [
                r'education',
                r'academic\s+background',
                r'educational\s+qualifications'
            ],
            'skills': [
                r'technical\s+skills',
                r'skills',
                r'core\s+competencies',
                r'technologies',
                r'expertise'
            ],
            'certifications': [
                r'certifications',
                r'certificates',
                r'professional\s+certifications'
            ],
            'projects': [
                r'projects',
                r'key\s+projects',
                r'notable\s+projects'
            ]
        }
    
    async def parse_resume(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse resume file and extract structured data"""
        try:
            if filename.lower().endswith('.pdf'):
                return await self._parse_pdf(file_content)
            elif filename.lower().endswith(('.docx', '.doc')):
                return await self._parse_docx(file_content)
            else:
                raise ValueError(f"Unsupported file format: {filename}")
        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {str(e)}")
            raise
    
    async def _parse_pdf(self, content: bytes) -> Dict[str, Any]:
        """Parse PDF resume using pdfplumber"""
        text_content = ""
        
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n"
        
        return await self._extract_structured_data(text_content)
    
    async def _parse_docx(self, content: bytes) -> Dict[str, Any]:
        """Parse DOCX resume using python-docx"""
        doc = docx.Document(BytesIO(content))
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Also extract from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"
        
        return await self._extract_structured_data(text_content)
    
    async def _extract_structured_data(self, text: str) -> Dict[str, Any]:
        """Extract structured data from raw text"""
        sections = self._identify_sections(text)
        
        return {
            "personal_info": self._extract_personal_info(text),
            "summary": sections.get("summary", ""),
            "experience": self._extract_experience(sections.get("experience", "")),
            "education": self._extract_education(sections.get("education", "")),
            "skills": self._extract_skills(sections.get("skills", "")),
            "certifications": self._extract_certifications(sections.get("certifications", "")),
            "projects": self._extract_projects(sections.get("projects", "")),
            "raw_text": text,
            "sections": [
                {
                    "type": section_type,
                    "content": content,
                    "suggestions": []
                }
                for section_type, content in sections.items()
            ]
        }
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify different resume sections"""
        lines = text.split('\n')
        sections = {}
        current_section = None
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            section_found = None
            for section_type, patterns in self.section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line_lower):
                        section_found = section_type
                        break
                if section_found:
                    break
            
            if section_found:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                current_section = section_found
                current_content = []
            elif current_section:
                current_content.append(line)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def _extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extract personal contact information"""
        info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            info['email'] = emails[0]
        
        # Phone pattern
        phone_pattern = r'(\+?1?[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})'
        phones = re.findall(phone_pattern, text)
        if phones:
            info['phone'] = ''.join(phones[0])
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            info['linkedin'] = linkedin.group()
        
        # GitHub pattern
        github_pattern = r'github\.com/[\w-]+'
        github = re.search(github_pattern, text, re.IGNORECASE)
        if github:
            info['github'] = github.group()
        
        return info
    
    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries"""
        experiences = []
        
        # Split by common separators (this is a simplified approach)
        # In a production system, you'd want more sophisticated parsing
        entries = re.split(r'\n\s*\n', text)
        
        for entry in entries:
            if entry.strip():
                experience = {
                    "raw_text": entry.strip(),
                    "company": "",
                    "position": "",
                    "duration": "",
                    "description": entry.strip()
                }
                experiences.append(experience)
        
        return experiences
    
    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries"""
        education = []
        
        entries = re.split(r'\n\s*\n', text)
        
        for entry in entries:
            if entry.strip():
                edu = {
                    "raw_text": entry.strip(),
                    "institution": "",
                    "degree": "",
                    "field": "",
                    "year": "",
                    "description": entry.strip()
                }
                education.append(edu)
        
        return education
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills list"""
        # Common skill separators
        skills = []
        
        # Split by common separators
        for separator in [',', '•', '·', '\n', '|']:
            if separator in text:
                skills = [skill.strip() for skill in text.split(separator)]
                break
        
        # Filter out empty strings and clean up
        skills = [skill for skill in skills if skill and len(skill) > 1]
        
        return skills
    
    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certifications = []
        
        # Split by lines or common separators
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) > 3:
                certifications.append(line)
        
        return certifications
    
    def _extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract project entries"""
        projects = []
        
        entries = re.split(r'\n\s*\n', text)
        
        for entry in entries:
            if entry.strip():
                project = {
                    "raw_text": entry.strip(),
                    "name": "",
                    "description": entry.strip(),
                    "technologies": [],
                    "url": ""
                }
                projects.append(project)
        
        return projects
